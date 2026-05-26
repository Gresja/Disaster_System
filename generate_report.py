from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Helper functions ──────────────────────────────────────────

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = run._element.get_or_add_rPr()
    s = shading.makeelement(qn('w:shd'), {})
    s.set(qn('w:fill'), 'F0F0F0')
    s.set(qn('w:val'), 'clear')
    shading.append(s)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Disaster Emergency Response\nCoordination System')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Project Report — Parallel Programming Course')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Language: Java\nConcurrency: java.util.concurrent\nVisualization: Real-Time Web Dashboard')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════

add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. System Architecture',
    '3. Parallel Programming Concepts Used',
    '4. Resource Requirements per Disaster Type',
    '5. Project Structure',
    '6. Execution Flow',
    '7. Web Dashboard',
    '8. Thread Safety Summary',
    '9. How to Run',
    '10. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

add_heading('1. Introduction', level=1)

add_heading('1.1 Problem Statement', level=2)
doc.add_paragraph(
    'When natural disasters occur — earthquakes, floods, forest fires, landslides, or storms — '
    'emergency centers receive many requests simultaneously. Injured people need ambulances, '
    'families need shelters, fire zones need rescue teams, roads may be blocked, and food/water '
    'supplies must be delivered.'
)
doc.add_paragraph(
    'If these requests are handled sequentially (one at a time), the response is dangerously slow. '
    'People die waiting in line while resources sit idle.'
)

add_heading('1.2 Proposed Solution', level=2)
doc.add_paragraph(
    'This project implements a Disaster Emergency Response Coordination System that uses '
    'parallel programming in Java to process multiple emergencies simultaneously. The system simulates:'
)
add_bullet('reporting emergencies at the same time (parallel producers)', bold_prefix='5 disaster zones ')
add_bullet('working on different emergencies simultaneously (parallel consumers)', bold_prefix='4 response handlers ')
add_bullet('(ambulances, shelters, etc.) that multiple handlers compete for safely', bold_prefix='Limited shared resources ')
add_bullet('to visualize the entire process', bold_prefix='A real-time web dashboard ')

add_heading('1.3 Technologies Used', level=2)
add_table(
    ['Technology', 'Purpose'],
    [
        ['Java 11+', 'Core programming language'],
        ['java.util.concurrent', 'Thread, Semaphore, ExecutorService, BlockingQueue, Atomic types'],
        ['com.sun.net.httpserver', 'Built-in HTTP server (no external dependencies)'],
        ['Server-Sent Events (SSE)', 'Real-time push from Java to browser'],
        ['HTML / CSS / JavaScript', 'Web dashboard'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════

add_heading('2. System Architecture', level=1)

add_heading('2.1 Architecture Diagram', level=2)
doc.add_paragraph(
    'The system follows the Producer-Consumer pattern with shared resources:'
)

arch_text = (
    'PRODUCER THREADS (DisasterZone × 5)\n'
    '  Zone-Alpha (Earthquake)    Zone-Beta (Flood)    Zone-Gamma (Forest Fire)\n'
    '  Zone-Delta (Landslide)     Zone-Epsilon (Storm)\n'
    '         │                        │                       │\n'
    '         └────────────────────────┴───────────────────────┘\n'
    '                                  │\n'
    '                   PriorityBlockingQueue<Emergency>\n'
    '                   (sorted by severity — worst first)\n'
    '                                  │\n'
    'CONSUMER THREADS (EmergencyHandler × 4 in ThreadPool)\n'
    '  Handler-1      Handler-2      Handler-3      Handler-4\n'
    '         │                        │                       │\n'
    '         └────────────────────────┴───────────────────────┘\n'
    '                                  │\n'
    '                   ResourceManager (Semaphore per type)\n'
    '                   Ambulance(8)  Shelter(5)  FireRescue(6)\n'
    '                   Food/Water(10)  RoadCrew(4)\n'
    '\n'
    'BACKGROUND: MonitorThread (daemon) + WebServer (port 8080)'
)
add_code(arch_text)

add_heading('2.2 Design Pattern: Producer-Consumer', level=2)
doc.add_paragraph(
    'The system follows the classic Producer-Consumer pattern:'
)
add_bullet('= 5 DisasterZone threads that generate emergencies', bold_prefix='Producers ')
add_bullet('= PriorityBlockingQueue that holds pending emergencies', bold_prefix='Shared Buffer ')
add_bullet('= 4 EmergencyHandler threads (in a thread pool) that process emergencies', bold_prefix='Consumers ')
doc.add_paragraph(
    'This pattern ensures that producers and consumers work independently and at their own speed, '
    'without data corruption or race conditions.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. PARALLEL PROGRAMMING CONCEPTS
# ═══════════════════════════════════════════════════════════════

add_heading('3. Parallel Programming Concepts Used', level=1)

# 3.1 Thread
add_heading('3.1 Thread (DisasterZone.java)', level=2)
doc.add_paragraph(
    'Each disaster zone is an independent Thread. All 5 zones run simultaneously, '
    'producing emergencies at random intervals.'
)
add_code(
    'public class DisasterZone extends Thread {\n'
    '    @Override\n'
    '    public void run() {\n'
    '        startSignal.await();   // wait for "go" signal\n'
    '        for (int i = 0; i < emergenciesCount; i++) {\n'
    '            Emergency emergency = new Emergency(zoneName, type, severity);\n'
    '            emergencyQueue.submit(emergency);\n'
    '            Thread.sleep(500 + random.nextInt(1000));\n'
    '        }\n'
    '    }\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('In a real disaster, multiple areas report emergencies at the same time. '
          'Each zone must operate independently — one zone should not wait for another to finish.')

# 3.2 CountDownLatch
add_heading('3.2 CountDownLatch (DisasterCoordinationSystem.java)', level=2)
doc.add_paragraph(
    'A CountDownLatch(1) ensures all 5 zone threads start at the exact same moment, '
    'simulating a simultaneous disaster onset.'
)
add_code(
    'CountDownLatch startLatch = new CountDownLatch(1);\n'
    '// ... create all zone threads (they call startSignal.await()) ...\n'
    'startLatch.countDown();  // releases ALL zones at once'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('Without this, zones would start at slightly different times depending on thread scheduling. '
          'The latch guarantees true simultaneous start.')

# 3.3 PriorityBlockingQueue
add_heading('3.3 PriorityBlockingQueue (EmergencyQueue.java)', level=2)
doc.add_paragraph('The emergency queue is a PriorityBlockingQueue<Emergency> that:')
add_bullet('most severe emergencies are handled first (Emergency implements Comparable)', bold_prefix='Sorts by severity — ')
add_bullet('multiple zones can put() and multiple handlers can take() simultaneously', bold_prefix='Is thread-safe — ')
add_bullet('handlers call take() which automatically waits if the queue is empty', bold_prefix='Blocks on empty — ')
add_code(
    'private final PriorityBlockingQueue<Emergency> queue\n'
    '    = new PriorityBlockingQueue<>();\n'
    '\n'
    'public void submit(Emergency emergency) {\n'
    '    queue.put(emergency);       // never blocks (unbounded)\n'
    '}\n'
    '\n'
    'public Emergency take() throws InterruptedException {\n'
    '    return queue.take();        // blocks if empty\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('Without a thread-safe queue, multiple threads writing/reading simultaneously would cause data corruption. '
          'The blocking behavior also prevents busy-waiting.')

# 3.4 ExecutorService
add_heading('3.4 ExecutorService / Thread Pool (ResourceDispatcher.java)', level=2)
doc.add_paragraph(
    'Instead of manually creating handler threads, we use Executors.newFixedThreadPool(4) '
    'which creates a pool of 4 reusable worker threads.'
)
add_code(
    'this.threadPool = Executors.newFixedThreadPool(handlerCount);\n'
    'for (EmergencyHandler handler : handlers) {\n'
    '    threadPool.submit(handler);   // each handler runs as a Runnable\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('Thread pools are more efficient than creating/destroying threads per task. '
          'The pool manages thread lifecycle, reuse, and shutdown.')

# 3.5 Semaphore
add_heading('3.5 Semaphore (ResourceManager.java)', level=2)
doc.add_paragraph(
    'Each resource type has its own counting Semaphore that controls concurrent access. '
    'For example, there are only 8 ambulances. If all 8 are deployed, the next handler '
    'that needs one must WAIT until another handler finishes and releases theirs.'
)
add_code(
    'semaphores.put(ResourceType.AMBULANCE, new Semaphore(8, true));\n'
    'semaphores.put(ResourceType.ROAD_CREW, new Semaphore(4, true));\n'
    '\n'
    'public boolean acquire(ResourceType type, int emergencyId)\n'
    '        throws InterruptedException {\n'
    '    semaphores.get(type).acquire();   // blocks if count == 0\n'
    '    return true;\n'
    '}\n'
    '\n'
    'public void release(ResourceType type, int emergencyId) {\n'
    '    semaphores.get(type).release();   // wakes up waiting thread\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('There are only 4 road crews, but multiple handlers might need one at the same time. '
          'The Semaphore prevents over-allocation and makes threads wait fairly.')

# 3.6 Atomic Variables
add_heading('3.6 Atomic Variables (AtomicInteger, AtomicLong, AtomicBoolean)', level=2)
doc.add_paragraph(
    'AtomicInteger and AtomicLong provide lock-free thread-safe counters. '
    'AtomicBoolean serves as a cooperative shutdown signal.'
)
add_code(
    'private final AtomicInteger totalReceived = new AtomicInteger(0);\n'
    'totalReceived.incrementAndGet();  // thread-safe increment\n'
    '\n'
    '// Shutdown signal checked by all 4 handlers:\n'
    'while (running.get() || emergencyQueue.getQueueSize() > 0) {\n'
    '    Emergency e = emergencyQueue.take();\n'
    '    handleEmergency(e);\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('Multiple threads update these counters simultaneously. Normal int++ would cause race conditions. '
          'AtomicBoolean ensures the shutdown flag is visible to all threads immediately.')

# 3.7 Daemon Thread
add_heading('3.7 Daemon Thread (MonitorThread.java)', level=2)
doc.add_paragraph(
    'The monitor thread is marked as setDaemon(true). It runs in the background, printing '
    'system status every 4 seconds and pushing updates to the web dashboard.'
)
add_code(
    'public MonitorThread(...) {\n'
    '    super("Monitor-Thread");\n'
    '    setDaemon(true);   // won\'t prevent JVM exit\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('Daemon threads are automatically terminated when all non-daemon threads finish. '
          'This prevents the monitor from keeping the program alive indefinitely.')

# 3.8 Synchronized
add_heading('3.8 Synchronized Logger (Logger.java)', level=2)
doc.add_paragraph(
    'All logging methods are synchronized to prevent garbled output when multiple threads '
    'print at the same time.'
)
add_code(
    'public static synchronized void info(String source, String msg) {\n'
    '    System.out.printf("[%s] [%s] %s%n", timestamp(), source, msg);\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run('Why needed: ')
run.bold = True
p.add_run('Without synchronization, concurrent System.out.println() calls from different threads '
          'would interleave characters, producing unreadable output.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. RESOURCE REQUIREMENTS
# ═══════════════════════════════════════════════════════════════

add_heading('4. Resource Requirements per Disaster Type', level=1)
doc.add_paragraph(
    'Each disaster type requires a specific set of resources. When a handler picks up an emergency, '
    'it must acquire ALL required resources before it can respond:'
)
add_table(
    ['Disaster Type', 'Required Resources'],
    [
        ['Earthquake', 'Ambulance + Shelter + Road Crew'],
        ['Flood', 'Food/Water Supply + Shelter + Road Crew'],
        ['Forest Fire', 'Fire Rescue Team + Ambulance'],
        ['Landslide', 'Road Crew + Ambulance + Food/Water Supply'],
        ['Storm', 'Shelter + Food/Water Supply + Ambulance'],
    ]
)

doc.add_paragraph()
add_heading('Available Resource Pool', level=2)
add_table(
    ['Resource', 'Total Units'],
    [
        ['Ambulance', '8'],
        ['Shelter', '5'],
        ['Fire Rescue Team', '6'],
        ['Food/Water Supply', '10'],
        ['Road Crew', '4'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Key observation: ')
run.bold = True
p.add_run('Road Crew has only 4 units but is needed by 3 disaster types (Earthquake, Flood, Landslide). '
          'This creates resource contention — handlers must wait when all road crews are deployed, '
          'demonstrating the Semaphore blocking behavior.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. PROJECT STRUCTURE
# ═══════════════════════════════════════════════════════════════

add_heading('5. Project Structure', level=1)
doc.add_paragraph('The project contains 14 Java source files organized into 5 packages, plus a web dashboard:')
add_table(
    ['File', 'Package', 'Purpose'],
    [
        ['DisasterCoordinationSystem.java', '(default)', 'Main entry point — orchestrates everything'],
        ['DisasterType.java', 'models', 'Enum: 5 disaster types with base severity'],
        ['ResourceType.java', 'models', 'Enum: 5 resource types with capacities'],
        ['Emergency.java', 'models', 'Emergency model, Comparable by severity'],
        ['EmergencyQueue.java', 'managers', 'PriorityBlockingQueue wrapper'],
        ['ResourceManager.java', 'managers', 'Semaphore-based resource pool'],
        ['StatisticsManager.java', 'managers', 'Atomic counters for response stats'],
        ['DisasterZone.java', 'threads', 'Producer thread (1 per zone)'],
        ['EmergencyHandler.java', 'threads', 'Consumer Runnable (4 in pool)'],
        ['ResourceDispatcher.java', 'threads', 'ExecutorService thread pool manager'],
        ['MonitorThread.java', 'threads', 'Background dashboard updater'],
        ['Logger.java', 'utils', 'Synchronized colored console logger'],
        ['EventBroadcaster.java', 'web', 'SSE event broadcaster (Singleton)'],
        ['WebServer.java', 'web', 'Built-in HTTP server for dashboard'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. EXECUTION FLOW
# ═══════════════════════════════════════════════════════════════

add_heading('6. Execution Flow (Step by Step)', level=1)

steps = [
    ('Step 1: System Initialization',
     'The main method creates the shared infrastructure: EmergencyQueue (shared priority queue), '
     'ResourceManager (initializes Semaphores), StatisticsManager (atomic counters), '
     'and WebServer (HTTP server on port 8080).'),
    ('Step 2: Start Background Threads',
     'MonitorThread starts as a daemon, printing status every 4 seconds. '
     'ResourceDispatcher creates a thread pool with 4 EmergencyHandler workers.'),
    ('Step 3: Create Zone Threads',
     '5 DisasterZone threads are created and started. Each one immediately blocks on '
     'startSignal.await(), waiting for the go signal.'),
    ('Step 4: Simultaneous Release',
     'startLatch.countDown() is called — all 5 zone threads unblock simultaneously '
     'and begin generating emergencies.'),
    ('Step 5: Parallel Processing',
     'Zones produce emergencies at random intervals (500–1500ms) and add them to the shared queue. '
     'Handlers pull the highest-severity emergency, acquire resources via Semaphore, '
     'simulate response, release resources, and loop. '
     'Monitor periodically reads shared state and updates both console and web dashboard.'),
    ('Step 6: Resource Contention',
     'When a resource runs out (e.g., all 4 Road Crews are deployed), the handler that needs one '
     'blocks on semaphore.acquire(). It automatically resumes when another handler calls semaphore.release().'),
    ('Step 7: Shutdown',
     'Once all zones finish and the queue is drained: systemRunning.set(false) signals handlers to stop, '
     'dispatcher.shutdown() calls threadPool.shutdown() and waits, final statistics are printed, '
     'and the web dashboard shows a completion banner.'),
]
for title, desc in steps:
    add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. WEB DASHBOARD
# ═══════════════════════════════════════════════════════════════

add_heading('7. Web Dashboard', level=1)
doc.add_paragraph(
    'The system includes a real-time web dashboard accessible at http://localhost:8080. '
    'It uses Server-Sent Events (SSE) — a persistent HTTP connection where the Java server '
    'pushes JSON events to the browser as they happen.'
)

add_heading('Dashboard Components', level=2)
add_table(
    ['Section', 'What it shows'],
    [
        ['Flow Diagram', '4-step pipeline: Disasters → Queue → Handlers → Resolved'],
        ['Disaster Zones', 'Status of each zone (Waiting / Active / Done) with emergency count'],
        ['Resource Bars', 'Real-time utilization of each resource type (green/yellow/red)'],
        ['Resource Requirements', 'What each disaster type needs from the resource pool'],
        ['Event Log', 'Chronological feed: REPORTED → HANDLING → RESOLVED for each emergency'],
    ]
)

add_heading('SSE Architecture', level=2)
doc.add_paragraph(
    'When the browser opens the dashboard, it creates an EventSource connection to /events. '
    'The Java server keeps this connection open and pushes JSON data whenever something happens '
    '(emergency submitted, resource acquired, emergency resolved, etc.). '
    'The browser JavaScript parses each event and updates the DOM in real time.'
)
add_code(
    'Browser:  GET /events\n'
    'Server:   Content-Type: text/event-stream\n'
    'Server:   data: {"type":"emergency_submitted","id":1,...}\n'
    'Server:   data: {"type":"resource_acquired","resource":"AMBULANCE",...}\n'
    'Server:   data: {"type":"emergency_resolved","id":1,"responseTimeMs":2500}\n'
    '          ... continuous stream until simulation ends ...'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. THREAD SAFETY SUMMARY
# ═══════════════════════════════════════════════════════════════

add_heading('8. Thread Safety Summary', level=1)
doc.add_paragraph(
    'Every piece of shared state in the system is protected by an appropriate concurrency mechanism:'
)
add_table(
    ['Shared Resource', 'Protection Mechanism', 'Why'],
    [
        ['Emergency queue', 'PriorityBlockingQueue', 'Multiple producers + consumers access simultaneously'],
        ['Resources (ambulances, etc.)', 'Semaphore (per type)', 'Limits concurrent access to finite resources'],
        ['Received/Resolved counters', 'AtomicInteger', 'Lock-free increment by many threads'],
        ['Response time accumulator', 'AtomicLong', 'Lock-free addition by many threads'],
        ['System running flag', 'AtomicBoolean', 'Visible to all threads immediately'],
        ['Console output', 'synchronized methods', 'Prevents interleaved printing'],
        ['SSE broadcast', 'synchronized + CopyOnWriteArrayList', 'Multiple threads emit events simultaneously'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. HOW TO RUN
# ═══════════════════════════════════════════════════════════════

add_heading('9. How to Run', level=1)

add_heading('Prerequisites', level=2)
add_bullet('Java 11 or higher installed')

add_heading('Steps', level=2)
add_code(
    'cd "Disaster System"\n'
    'chmod +x run.sh\n'
    './run.sh'
)
doc.add_paragraph(
    'The script compiles all source files and runs the system. '
    'The browser opens automatically to http://localhost:8080. '
    'Press Ctrl+C to stop.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. CONCLUSION
# ═══════════════════════════════════════════════════════════════

add_heading('10. Conclusion', level=1)
doc.add_paragraph(
    'This project demonstrates key parallel programming concepts through a realistic '
    'disaster response simulation:'
)
add_bullet('running simultaneously (zones + handlers + monitor)', bold_prefix='Multiple threads ')
add_bullet('using PriorityBlockingQueue, Semaphore, and Atomic types', bold_prefix='Thread-safe data sharing ')
add_bullet('with CountDownLatch (simultaneous start) and synchronized (logging)', bold_prefix='Synchronization ')
add_bullet('with ExecutorService for efficient thread management', bold_prefix='Thread pooling ')
add_bullet('where threads must wait for limited shared resources', bold_prefix='Resource contention ')
add_bullet('using cooperative AtomicBoolean signaling', bold_prefix='Graceful shutdown ')

doc.add_paragraph()
doc.add_paragraph(
    'The system processes 12 emergencies from 5 zones using 4 parallel handlers, achieving '
    'significantly faster response times than a sequential approach would allow. The web dashboard '
    'makes the parallelism visible in real time, showing how multiple threads interact, compete '
    'for resources, and resolve emergencies simultaneously.'
)

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════

output_path = '/Users/apple/Desktop/Disaster System/Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
